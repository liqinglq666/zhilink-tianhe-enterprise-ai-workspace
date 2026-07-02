# -*- coding: utf-8 -*-
"""报告导出工具。

用于把各模块生成结果导出为 Markdown / TXT / Word DOCX。

重点优化：
1. Markdown 表格转真实 Word 表格，不再显示 |---| 原始符号；
2. 支持无首尾竖线的表格行；
3. 支持中文全角竖线 ｜；
4. 支持 **加粗**、`代码`、列表、复选框；
5. 宽表格自动使用横向页面；
6. 合同审阅、风险清单、实施计划等内容导出更像正式报告。
"""

from __future__ import annotations

import re
from datetime import datetime
from io import BytesIO
from typing import Dict, List, Sequence, Tuple

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    DOCX_AVAILABLE = True
except Exception:  # noqa: BLE001
    DOCX_AVAILABLE = False

from .constants import APP_FULL_NAME, LEGAL_DISCLAIMER


KNOWN_SECTION_TITLES = {
    "一句话结论",
    "风险总览",
    "重点风险清单",
    "建议补充条款",
    "签约前检查清单",
    "免责声明",
    "企业画像",
    "主要运营痛点",
    "推荐使用模块",
    "后续动作清单",
    "会议摘要",
    "关键决策",
    "待办事项",
    "风险提醒",
    "政策方向",
    "材料清单",
    "合作方案框架",
    "实施路径",
    "数据边界",
    "复核机制",
    "评估指标",
    "下一步动作",
    "下次会议议题",
}


def build_markdown_report(results: Dict[str, str]) -> str:
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    parts = [
        f"# {APP_FULL_NAME}运营报告",
        f"生成时间：{now}",
        "",
        f"> {LEGAL_DISCLAIMER}",
        "",
    ]

    for title, content in results.items():
        if content:
            parts.append(f"## {title}")
            parts.append(content.strip())
            parts.append("")

    return "\n".join(parts).strip()


def build_txt_bytes(results: Dict[str, str]) -> bytes:
    return build_markdown_report(results).encode("utf-8")


def _set_east_asian_font(run, font_name: str = "微软雅黑") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_format(paragraph, *, before=0, after=6, line_spacing=1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")

    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)

        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str = "D0D5DD", size: str = "6") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _clean_inline_text(text: str) -> str:
    return (
        str(text)
        .replace("&nbsp;", " ")
        .replace("<br>", "\n")
        .replace("<br/>", "\n")
        .replace("<br />", "\n")
        .replace("｜", "|")
        .strip()
    )


def _add_inline_runs(paragraph, text: str, *, base_size: float = 10.5, color: str | None = None) -> None:
    """把简单 Markdown 行内语法转换为 Word Run。

    支持：
    - **加粗**
    - __加粗__
    - `代码`
    """

    text = _clean_inline_text(text)
    if not text:
        return

    pattern = re.compile(r"(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`)")
    pos = 0

    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            _set_east_asian_font(run)
            run.font.size = Pt(base_size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)

        token = match.group(0)
        run_text = token
        bold = False
        code = False

        if token.startswith("**") and token.endswith("**"):
            run_text = token[2:-2]
            bold = True
        elif token.startswith("__") and token.endswith("__"):
            run_text = token[2:-2]
            bold = True
        elif token.startswith("`") and token.endswith("`"):
            run_text = token[1:-1]
            code = True

        run = paragraph.add_run(run_text)
        _set_east_asian_font(run, "Consolas" if code else "微软雅黑")
        run.font.size = Pt(base_size if not code else max(base_size - 0.5, 8))
        run.bold = bold

        if code:
            run.font.color.rgb = RGBColor(30, 64, 175)
        elif color:
            run.font.color.rgb = RGBColor.from_string(color)

        pos = match.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_east_asian_font(run)
        run.font.size = Pt(base_size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _normalize_table_line(line: str) -> str:
    line = _clean_inline_text(line)

    # 有些模型会把表格行前面加项目符号，这里统一去掉。
    line = re.sub(r"^[•\-*+]\s*(?=\|)", "", line)

    # 有些生成结果可能包含 Word 复制出来的换行符号，提前清理。
    line = line.replace("\r", "").replace("\n", "")

    return line.strip()


def _is_table_separator_cells(cells: Sequence[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c.strip()) for c in cells)


def _split_table_row(line: str) -> List[str]:
    line = _normalize_table_line(line)

    if line.startswith("|"):
        line = line[1:]

    if line.endswith("|"):
        line = line[:-1]

    return [cell.strip() for cell in line.split("|")]


def _is_probable_table_line(line: str) -> bool:
    """判断一行是否像 Markdown 表格。

    比旧版更宽松：
    - 支持 | A | B |
    - 支持 A | B | C
    - 支持中文全角竖线 ｜
    """

    line = _normalize_table_line(line)

    if not line:
        return False

    # 至少要有两个竖线，才认为可能是多列表格。
    if line.count("|") < 2:
        return False

    cells = _split_table_row(line)

    # 至少两列，而且不能全为空。
    if len(cells) < 2:
        return False

    if not any(cell.strip() for cell in cells):
        return False

    return True


def _collect_table(lines: Sequence[str], start: int) -> Tuple[List[List[str]], int]:
    """从 start 开始收集连续表格行。

    支持标准 Markdown 表格，也支持模型生成的非严格表格。
    """

    raw_rows: List[List[str]] = []
    i = start

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            # 空行代表表格结束。
            break

        if not _is_probable_table_line(line):
            break

        cells = _split_table_row(line)

        # 跳过 |---|---| 分隔行。
        if not _is_table_separator_cells(cells):
            raw_rows.append(cells)

        i += 1

    if not raw_rows:
        return [], i

    max_cols = max(len(row) for row in raw_rows)
    rows = [row + [""] * (max_cols - len(row)) for row in raw_rows]

    return rows, i


def _add_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return

    col_count = len(rows[0])
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.autofit = True

    _set_table_borders(table)

    for r_idx, row_data in enumerate(rows):
        row = table.add_row()

        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(cell)

            if r_idx == 0:
                _set_cell_shading(cell, "EAF2FF")
            elif r_idx % 2 == 0:
                _set_cell_shading(cell, "FAFBFF")

            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_paragraph_format(paragraph, before=0, after=0, line_spacing=1.05)

            # 宽表格自动缩小字号。
            if col_count >= 6:
                font_size = 7.8
            elif col_count >= 5:
                font_size = 8.4
            elif col_count >= 4:
                font_size = 9
            else:
                font_size = 9.5

            _add_inline_runs(paragraph, cell_text, base_size=font_size)

            for run in paragraph.runs:
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(17, 24, 39)

    doc.add_paragraph()


def _looks_like_plain_section_title(text: str) -> bool:
    stripped = _clean_inline_text(text).strip("：:")

    if stripped in KNOWN_SECTION_TITLES:
        return True

    if len(stripped) <= 14 and re.fullmatch(r"[\u4e00-\u9fffA-Za-z0-9·（）() ]+", stripped):
        title_keywords = (
            "清单",
            "总览",
            "结论",
            "说明",
            "建议",
            "机制",
            "框架",
            "风险",
            "摘要",
            "指标",
            "免责声明",
            "议题",
            "动作",
        )
        return any(k in stripped for k in title_keywords)

    return False


def _add_heading(doc: Document, text: str, level: int) -> None:
    heading = doc.add_heading("", level=level)
    _set_paragraph_format(heading, before=8 if level <= 2 else 6, after=6, line_spacing=1.1)

    run = heading.add_run(_clean_inline_text(text))
    _set_east_asian_font(run)
    run.bold = True

    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(17, 24, 39)
    elif level == 2:
        run.font.size = Pt(13.5)
        run.font.color.rgb = RGBColor(29, 78, 216)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(52, 64, 84)


def _add_quote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.35)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)

    _add_inline_runs(paragraph, text, base_size=9.5, color="667085")


def _add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    _set_paragraph_format(paragraph, before=0, after=6, line_spacing=1.18)
    _add_inline_runs(paragraph, text, base_size=10.5)


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    _set_paragraph_format(paragraph, before=0, after=4, line_spacing=1.12)
    _add_inline_runs(paragraph, text, base_size=10.2)


def _add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    _set_paragraph_format(paragraph, before=0, after=4, line_spacing=1.12)
    _add_inline_runs(paragraph, text, base_size=10.2)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]

    # 横向页面更适合风险清单、合同审阅等宽表格。
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    styles = doc.styles

    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        if style_name in styles:
            styles[style_name].font.name = "微软雅黑"
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)

    run = title.add_run("智链天河：企业运营 AI 工作台")
    _set_east_asian_font(run)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(17, 24, 39)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)

    run = subtitle.add_run("运营材料生成与风险提示报告")
    _set_east_asian_font(run)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(102, 112, 133)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)

    run = p.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    _set_east_asian_font(run)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(102, 112, 133)

    disclaimer = doc.add_paragraph()
    disclaimer.paragraph_format.left_indent = Cm(0.2)
    disclaimer.paragraph_format.right_indent = Cm(0.2)
    disclaimer.paragraph_format.space_after = Pt(12)

    _add_inline_runs(disclaimer, f"说明：{LEGAL_DISCLAIMER}", base_size=9.2, color="667085")


def _add_markdown_content(doc: Document, content: str) -> None:
    lines = content.splitlines()
    i = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if not line:
            i += 1
            continue

        # 关键：优先识别表格，避免表格被当成普通段落输出。
        if _is_probable_table_line(line):
            table_rows, next_i = _collect_table(lines, i)

            if table_rows:
                _add_table(doc, table_rows)
                i = next_i
                continue

        if re.fullmatch(r"-{3,}|_{3,}|\*{3,}", line):
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("# "):
            _add_heading(doc, line[2:].strip(), 1)

        elif line.startswith("## "):
            _add_heading(doc, line[3:].strip(), 2)

        elif line.startswith("### "):
            _add_heading(doc, line[4:].strip(), 3)

        elif line.startswith("> "):
            _add_quote(doc, line[2:].strip())

        elif re.match(r"^[-*+]\s+\[[ xX]\]\s+", line):
            checked = bool(re.match(r"^[-*+]\s+\[[xX]\]\s+", line))
            text = re.sub(r"^[-*+]\s+\[[ xX]\]\s+", "", line).strip()
            _add_bullet(doc, ("☑ " if checked else "☐ ") + text)

        elif line.startswith("•"):
            text = line[1:].strip()

            # 如果项目符号后面其实是一个表格行，则也转成表格。
            if _is_probable_table_line(text):
                table_rows, next_i = _collect_table([text] + list(lines[i + 1:]), 0)
                if table_rows:
                    _add_table(doc, table_rows)
                    i = i + next_i
                    continue

            _add_bullet(doc, text)

        elif re.match(r"^[-*+]\s+", line):
            text = re.sub(r"^[-*+]\s+", "", line).strip()

            if _is_probable_table_line(text):
                table_rows, next_i = _collect_table([text] + list(lines[i + 1:]), 0)
                if table_rows:
                    _add_table(doc, table_rows)
                    i = i + next_i
                    continue

            _add_bullet(doc, text)

        elif re.match(r"^\d+[.)、]\s+", line):
            _add_numbered(doc, re.sub(r"^\d+[.)、]\s+", "", line).strip())

        elif _looks_like_plain_section_title(line):
            _add_heading(doc, line.rstrip("：:"), 2)

        else:
            _add_paragraph(doc, line)

        i += 1


def build_docx_bytes(results: Dict[str, str]) -> bytes:
    if not DOCX_AVAILABLE:
        raise RuntimeError("未安装 python-docx，无法导出 DOCX。")

    doc = Document()
    _configure_document(doc)
    _add_cover(doc)

    for section_title, content in results.items():
        if not content:
            continue

        _add_heading(doc, section_title, 1)
        _add_markdown_content(doc, content.strip())

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("智链天河 · 企业运营 AI 工作台")
    _set_east_asian_font(run)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(152, 162, 179)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()