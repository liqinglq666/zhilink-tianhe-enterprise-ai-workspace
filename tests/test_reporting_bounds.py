from __future__ import annotations

from io import BytesIO

from docx import Document

from zhilian_tianhe_agent.reporting import (
    MAX_DOCX_TABLE_COLUMNS,
    _collect_table,
    build_docx_bytes,
)


def _wide_table(column_count: int = 30) -> list[str]:
    header = "| " + " | ".join(f"列{index}" for index in range(column_count)) + " |"
    separator = "| " + " | ".join("---" for _ in range(column_count)) + " |"
    row = "| " + " | ".join(f"值{index}" for index in range(column_count)) + " |"
    return [header, separator, row]


def test_collect_table_caps_extreme_markdown_width() -> None:
    rows, next_index = _collect_table(_wide_table(), 0)

    assert next_index == 3
    assert len(rows) == 2
    assert all(len(row) == MAX_DOCX_TABLE_COLUMNS for row in rows)
    assert rows[0][-1] == "列19"
    assert rows[1][-1] == "值19"


def test_docx_export_never_builds_table_wider_than_limit() -> None:
    data = build_docx_bytes({"测试模块": "\n".join(_wide_table())})
    document = Document(BytesIO(data))

    assert document.tables
    assert len(document.tables[0].columns) == MAX_DOCX_TABLE_COLUMNS
    assert document.tables[0].cell(0, MAX_DOCX_TABLE_COLUMNS - 1).text == "列19"
    assert document.tables[0].cell(1, MAX_DOCX_TABLE_COLUMNS - 1).text == "值19"
